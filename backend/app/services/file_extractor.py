"""
File text extraction utilities.

Supported formats:
- .txt  → Direct read (UTF-8 with fallback)
- .pdf  → PyMuPDF (fitz)
- .docx → python-docx

Design choices:
- Minimal dependencies (no OCR, no complex parsing)
- Graceful fallbacks for encoding issues
- Returns clean text ready for chunking
"""

import io
from pathlib import Path
from typing import Optional

# File size limit: 10MB (reasonable for demo, prevents memory issues)
MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024

# Supported extensions
SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


class FileExtractionError(Exception):
    """Raised when file extraction fails."""
    pass


class UnsupportedFileTypeError(FileExtractionError):
    """Raised for unsupported file types."""
    pass


class FileTooLargeError(FileExtractionError):
    """Raised when file exceeds size limit."""
    pass


def validate_file(filename: str, file_size: int) -> str:
    """
    Validate file type and size.
    
    Returns:
        File extension (lowercase)
        
    Raises:
        UnsupportedFileTypeError: If file type not supported
        FileTooLargeError: If file exceeds size limit
    """
    ext = Path(filename).suffix.lower()
    
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )
    
    if file_size > MAX_FILE_SIZE_BYTES:
        raise FileTooLargeError(
            f"File too large: {file_size / 1024 / 1024:.1f}MB. Max: {MAX_FILE_SIZE_MB}MB"
        )
    
    return ext


def extract_text_from_txt(content: bytes) -> str:
    """Extract text from .txt file with encoding fallback."""
    # Try UTF-8 first, then latin-1 (accepts any byte)
    for encoding in ["utf-8", "latin-1"]:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    
    raise FileExtractionError("Could not decode text file")


def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from PDF using PyMuPDF.
    
    Note: Only extracts text, no OCR for scanned documents.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise FileExtractionError(
            "PyMuPDF not installed. Run: pip install PyMuPDF"
        )
    
    try:
        # Open PDF from bytes
        doc = fitz.open(stream=content, filetype="pdf")
        
        text_parts = []
        for page_num, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)
        
        doc.close()
        
        if not text_parts:
            raise FileExtractionError(
                "No text found in PDF. May be scanned/image-based."
            )
        
        return "\n\n".join(text_parts)
    
    except Exception as e:
        if isinstance(e, FileExtractionError):
            raise
        raise FileExtractionError(f"PDF extraction failed: {str(e)}")


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise FileExtractionError(
            "python-docx not installed. Run: pip install python-docx"
        )
    
    try:
        # Open DOCX from bytes
        doc = Document(io.BytesIO(content))
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        
        if not text_parts:
            raise FileExtractionError("No text found in DOCX file.")
        
        return "\n\n".join(text_parts)
    
    except Exception as e:
        if isinstance(e, FileExtractionError):
            raise
        raise FileExtractionError(f"DOCX extraction failed: {str(e)}")


def extract_text(filename: str, content: bytes) -> str:
    """
    Extract text from file based on extension.
    
    Args:
        filename: Original filename (used to detect type)
        content: Raw file bytes
        
    Returns:
        Extracted text string
        
    Raises:
        FileExtractionError: If extraction fails
    """
    ext = Path(filename).suffix.lower()
    
    extractors = {
        ".txt": extract_text_from_txt,
        ".pdf": extract_text_from_pdf,
        ".docx": extract_text_from_docx,
    }
    
    extractor = extractors.get(ext)
    if not extractor:
        raise UnsupportedFileTypeError(f"No extractor for: {ext}")
    
    text = extractor(content)
    
    # Clean up text
    text = text.strip()
    
    # Remove excessive whitespace
    import re
    text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 newlines
    text = re.sub(r' {2,}', ' ', text)       # Max 1 space
    
    return text
